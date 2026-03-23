from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from PIL import Image
from pathlib import Path
from core.config import logger, IMAGE_EXTENSIONS

class DocumentService:
    @staticmethod
    def generate_evidence_doc(template_path, image_dir, output_path, data, test_data, is_devops):
        try:
            doc = Document(template_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX template: {e}")
            raise IOError(f"Failed to open template document: {str(e)}")

        DocumentService._insert_images(doc, image_dir, "Evidences", is_devops, test_data)
        DocumentService._replace_placeholders(doc, data)
        
        try:
            doc.save(output_path)
            logger.info(f"Document saved successfully at {output_path}")
        except PermissionError:
            logger.error("Permission denied when saving document.")
            raise IOError("Permission denied. Ensure the target Word document is closed.")

    @staticmethod
    def _is_valid_image(img_path):
        try:
            with Image.open(img_path) as img:
                img.verify()
            return True
        except Exception as e:
            logger.warning(f"Invalid or corrupted image ignored: {img_path} - {e}")
            return False

    @staticmethod
    def _insert_images(doc, image_dir, marker_text, is_devops_active, test_data):
        insert_paragraph = next((p for p in doc.paragraphs if marker_text in p.text), None)
        if not insert_paragraph:
            logger.warning(f"Marker '{marker_text}' not found in template.")
            return 

        parent = insert_paragraph._element.getparent()
        current_index = list(parent).index(insert_paragraph._element) + 1 
        
        p_space_initial = doc.add_paragraph("\n")
        parent.insert(current_index, p_space_initial._element)
        current_index += 1

        pre_requisite = test_data.get('pre_requisite') if is_devops_active else None
        steps = test_data.get('steps', []) if is_devops_active else []

        if pre_requisite:
            p_prereq = doc.add_paragraph()
            p_prereq.add_run(f"Pré-requisito: {pre_requisite}").bold = True
            parent.insert(current_index, p_prereq._element)
            current_index += 1
            
            p_space = doc.add_paragraph("\n")
            parent.insert(current_index, p_space._element)
            current_index += 1

        image_files = sorted(
            [f for f in Path(image_dir).iterdir() if f.suffix.lower() in IMAGE_EXTENSIONS and DocumentService._is_valid_image(f)], 
            key=lambda x: x.stat().st_mtime
        )
        
        limit = max(len(image_files), len(steps)) if is_devops_active else len(image_files)

        if limit == 0:
            p = doc.add_paragraph("Nenhuma imagem ou passo de teste encontrado.")
            parent.insert(current_index, p._element)
            current_index += 1
            logger.info("No images or steps found to insert.")
        else:
            for i in range(limit):
                step_info = steps[i] if is_devops_active and i < len(steps) else None
                img_path = image_files[i] if i < len(image_files) else None

                if step_info:
                    p_step = doc.add_paragraph()
                    p_step.add_run(f"STEP: {step_info.get('step', '-')}").bold = True
                    parent.insert(current_index, p_step._element)
                    current_index += 1

                    p_expected = doc.add_paragraph(f"OUTCOME: {step_info.get('expected', '-')}")
                    parent.insert(current_index, p_expected._element)
                    current_index += 1

                p_content = doc.add_paragraph()
                p_content.alignment = 1 
                
                if img_path:
                    p_content.add_run().add_picture(str(img_path), width=Inches(6.0))
                elif step_info:
                    p_content.add_run("[Image]").bold = True
                
                parent.insert(current_index, p_content._element)
                current_index += 1
                
                p_space = doc.add_paragraph("\n")
                parent.insert(current_index, p_space._element)
                current_index += 1

        p_break = doc.add_paragraph()
        p_break.add_run().add_break(WD_BREAK.PAGE)
        parent.insert(current_index, p_break._element)
        current_index += 1
        
        p_space_final = doc.add_paragraph("\n")
        parent.insert(current_index, p_space_final._element)

    @staticmethod
    def _replace_placeholders(doc, values):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        DocumentService._replace_in_paragraph(p, values)
        for p in doc.paragraphs:
            DocumentService._replace_in_paragraph(p, values)

    @staticmethod
    def _replace_in_paragraph(paragraph, values):
        for placeholder, value in values.items():
            if placeholder in paragraph.text:
                full_text = "".join([run.text for run in paragraph.runs])
                if placeholder in full_text:
                    new_text = full_text.replace(placeholder, str(value))
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.clear()