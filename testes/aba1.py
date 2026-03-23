import base64
import requests
import tkinter as tk
from tkinter import Toplevel, ttk, messagebox
from docx import Document
from docx.shared import Pt, Inches
from pathlib import Path
from base import decodificar_chave

# URL da organização no Azure DevOps
ORGANIZATION_URL = 'https://dev.azure.com/ProjetoTCC2'

def desativar_selecao(ambiente_dropdown, project_dropdown):
    """Limpa a seleção dos dropdowns para evitar inconsistências."""
    ambiente_dropdown.set('')
    project_dropdown.set('')

def cadastrar_renovar_chave(parent_window):
    """Abre uma janela para o usuário inserir e salvar a chave de acesso do Azure DevOps."""
    dialog = Toplevel(parent_window)
    dialog.title("Chave do Azure DevOps")
    dialog.geometry("350x150")
    dialog.transient(parent_window)
    dialog.grab_set() # Torna a janela modal
    dialog.resizable(False, False)

    ttk.Label(dialog, text="Insira seu Personal Access Token (PAT):").pack(pady=10)
    entry = ttk.Entry(dialog, width=50, show="*")
    entry.pack(pady=5, padx=10, fill='x')

    def submit():
        chave_devops = entry.get().strip()
        if not chave_devops:
            messagebox.showwarning("Erro", "A chave não pode estar vazia.", parent=dialog)
            return
        try:
            chave_codificada = base64.b64encode(chave_devops.encode('utf-8')).decode('utf-8')
            with open("chave_devops.txt", "w") as file:
                file.write(chave_codificada)
            messagebox.showinfo("Sucesso", "Chave cadastrada com sucesso.", parent=dialog)
            dialog.destroy()
        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao salvar a chave: {e}", parent=dialog)

    ttk.Button(dialog, text="Salvar Chave", command=submit, style="Accent.TButton").pack(pady=10)
    dialog.focus_set()

def on_checkbox_clicked(checkbox_var):
    """Valida a chave do Azure DevOps quando a checkbox de integração é marcada."""
    if not checkbox_var.get():
        return False # Não faz nada se a caixa for desmarcada

    chave_decodificada = decodificar_chave()
    if not chave_decodificada:
        messagebox.showerror("Chave Não Encontrada", "Nenhuma chave do Azure DevOps foi encontrada. Por favor, cadastre sua chave para usar a integração.")
        checkbox_var.set(False)
        return False

    if not validar_credenciais(ORGANIZATION_URL, chave_decodificada):
        messagebox.showerror("Credenciais Inválidas", "A chave fornecida é inválida ou expirou. Verifique seu token ou cadastre um novo.")
        checkbox_var.set(False)
        return False
    
    return True

def validar_credenciais(organization_url, pat):
    """Valida as credenciais fazendo uma chamada simples à API do Azure DevOps."""
    try:
        url = f"{organization_url}/_apis/connectionData?api-version=6.0"
        headers = {"Authorization": f"Basic {base64.b64encode(f':{pat}'.encode('ascii')).decode('ascii')}"}
        response = requests.get(url, headers=headers, timeout=10)
        return response.status_code == 200
    except requests.exceptions.RequestException:
        return False

def get_devops_projects(pat):
    """Obtém a lista de projetos do Azure DevOps."""
    try:
        url = f"{ORGANIZATION_URL}/_apis/projects?api-version=6.0"
        headers = {"Authorization": f"Basic {base64.b64encode(f':{pat}'.encode()).decode()}"}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        return [project['name'] for project in response.json().get("value", [])]
    except requests.exceptions.RequestException as e:
        messagebox.showerror("Erro de Rede", f"Erro ao buscar projetos: {e}")
        return []

def update_project_dropdown(checkbox_var, project_dropdown):
    """Atualiza o dropdown com a lista de projetos do Azure DevOps."""
    if checkbox_var.get():
        chave_decodificada = decodificar_chave()
        if chave_decodificada:
            projects = get_devops_projects(chave_decodificada)
            project_dropdown.configure(values=projects)
            if projects:
                project_dropdown.set(projects[0]) # Seleciona o primeiro por padrão
    else:
        project_dropdown.set('')
        project_dropdown.configure(values=[])
        
def criar_arquivo(doc_path, img_dir, ct_id, tester_name, ambiente, perfil, bugs, projeto, checkbox_var):
    """Função principal que gera o documento de evidências."""
    # Validações Iniciais
    if not all([doc_path, img_dir, ct_id, tester_name]):
        messagebox.showwarning("Campos Obrigatórios", "Preencha todos os campos: Arquivo padrão, Diretório de imagens, ID do cenário e Nome do Tester.")
        return

    # Processamento do Documento
    try:
        documento = Document(doc_path)
        
        # Inserção de Imagens
        criar_arquivos_com_imagens(img_dir, documento, "Evidências")
        
        # Substituição de Placeholders
        valores = {
            '[Nome do Tester]': tester_name,
            '[Numero do CT]': ct_id,
            '[US]': '', # Valor padrão para US
            '[Ambiente]': ambiente,
            '[Perfil]': perfil,
            '[Bugs]': bugs if bugs else "Nenhum",
            '[Resultado]': '☐Passed  ☒Failed' if bugs else '☒Passed  ☐Failed'
        }
        
        substituir_placeholders(documento, valores)

        # Salva o documento
        nome_arquivo = f"{Path(doc_path).stem}_{ct_id}.docx"
        novo_arquivo_path = str(Path(img_dir) / nome_arquivo)
        documento.save(novo_arquivo_path)

        messagebox.showinfo("Sucesso", f"Documento criado com sucesso em:\n{novo_arquivo_path}")

    except FileNotFoundError:
        messagebox.showerror("Erro", f"Arquivo padrão não encontrado em: {doc_path}")
    except Exception as e:
        messagebox.showerror("Erro Inesperado", f"Ocorreu um erro ao gerar o documento: {e}")


def criar_arquivos_com_imagens(diretorio_raiz, documento, placeholder_text):
    """Insere imagens de um diretório em um documento Word, substituindo um placeholder."""
    diretorio = Path(diretorio_raiz)
    if not diretorio.is_dir():
        messagebox.showwarning("Aviso", f"O diretório de imagens '{diretorio_raiz}' não foi encontrado.")
        return

    # Localiza o parágrafo que contém o placeholder
    placeholder_paragraph = None
    for p in documento.paragraphs:
        if placeholder_text in p.text:
            placeholder_paragraph = p
            break
    
    if not placeholder_paragraph:
        messagebox.showwarning("Aviso", f"O placeholder '{placeholder_text}' não foi encontrado no documento.")
        return

    # Limpa o texto do parágrafo placeholder
    placeholder_paragraph.text = ""
    
    extensoes_imagem = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
    imagens = sorted(
        [f for f in diretorio.iterdir() if f.is_file() and f.suffix.lower() in extensoes_imagem],
        key=lambda x: x.stat().st_mtime
    )

    if not imagens:
        placeholder_paragraph.add_run("Nenhuma imagem de evidência foi encontrada no diretório.").italic = True
        return

    # Insere cada imagem em um novo parágrafo após o placeholder
    for arquivo_img in imagens:
        try:
            # A imagem é adicionada diretamente ao parágrafo do placeholder, que agora está vazio
            placeholder_paragraph.add_run().add_picture(str(arquivo_img), width=Inches(6.0))
            # Adiciona um novo parágrafo para a próxima imagem, para que não fiquem lado a lado
            placeholder_paragraph = documento.add_paragraph()
        except Exception as e:
            messagebox.showerror("Erro de Imagem", f"Erro ao inserir a imagem {arquivo_img.name}: {e}")

def substituir_placeholders(documento, valores):
    """Substitui placeholders de texto em parágrafos e tabelas do documento."""
    # Substituir em parágrafos
    for p in documento.paragraphs:
        for key, value in valores.items():
            if key in p.text:
                # Substituir mantendo o estilo é complexo. Uma abordagem mais simples:
                # Armazenar runs, substituir texto e reaplicar formatação
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, value)
                        inline[i].text = text
                        inline[i].font.name = 'EYInterstate Light'
                        inline[i].font.size = Pt(14)
    
    # Substituir em tabelas
    for table in documento.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in valores.items():
                         if key in p.text:
                            inline = p.runs
                            for i in range(len(inline)):
                                if key in inline[i].text:
                                    text = inline[i].text.replace(key, value)
                                    inline[i].text = text
                                    inline[i].font.name = 'EYInterstate Light'
                                    inline[i].font.size = Pt(14)