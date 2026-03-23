import tkinter as tk
from tkinter import ttk, messagebox, filedialog, Toplevel
from ttkthemes import ThemedTk
from PIL import Image, ImageTk
from threading import Thread
from pathlib import Path
import base64
import requests
from docx import Document
from docx.shared import Inches, Pt

# --- Constantes ---
ORGANIZATION_URL = 'https://dev.azure.com/ProjetoTCCFinal'
KEY_FILE = "chave_devops.txt"
ICON_FILE = "ey.ico"
KEY_ICON_FILE = "chave.png"
IMAGE_EXTENSIONS = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')

# Dicionários para placeholders e labels da UI
UI_LABELS = {
    "file": "Arquivo padrão",
    "img_dir": "Diretório de Imagens",
    "tester": "Nome do Tester",
    "test_id": "ID/Nome do cenário de teste",
    "env": "Ambiente",
    "product": "Produto",
    "profile": "Perfil",
    "bugs": "Bugs"
}

PLACEHOLDERS = {
    "tester": '[Nome do Tester]',
    "test_id": '[Numero do CT]',
    "us": '[US]',
    "env": '[Ambiente]',
    "profile": '[Perfil]',
    "bugs": '[Bugs]',
    "result": '[Resultado]'
}

class DevOpsService:
    """Encapsula toda a comunicação com a API do Azure DevOps."""

    def __init__(self, personal_access_token):
        self.pat = personal_access_token
        self.auth_headers = {
            'Content-Type': 'application/json-patch+json',
            'Authorization': f'Basic {base64.b64encode(f":{self.pat}".encode()).decode()}'
        }

    def get_projects(self):
        """Busca a lista de projetos no Azure DevOps."""
        url = f"{ORGANIZATION_URL}/_apis/projects?api-version=6.1-preview.4"
        try:
            response = requests.get(url, headers={'Authorization': self.auth_headers['Authorization']})
            response.raise_for_status()
            return [project['name'] for project in response.json().get("value", [])]
        except requests.RequestException as e:
            messagebox.showerror("Erro de API", f"Falha ao obter projetos: {e}")
            return []

    def get_test_case_relations(self, test_case_id):
        """Busca a User Story (US) relacionada a um caso de teste."""
        url = f"{ORGANIZATION_URL}/_apis/wit/workitems/{test_case_id}?$expand=relations&api-version=6.0"
        try:
            response = requests.get(url, headers={'Authorization': self.auth_headers['Authorization']})
            response.raise_for_status()
            work_item = response.json()
            
            test_case_title = work_item['fields'].get('System.Title', test_case_id)
            us_title = ""

            if 'relations' in work_item:
                for rel in work_item['relations']:
                    if rel['rel'] == 'Microsoft.VSTS.Common.TestedBy-Reverse':
                        us_url = rel['url']
                        us_response = requests.get(us_url, headers={'Authorization': self.auth_headers['Authorization']})
                        us_response.raise_for_status()
                        us_item = us_response.json()
                        if us_item['fields']['System.WorkItemType'] in ['Product Backlog Item', 'User Story']:
                            us_title = us_item['fields']['System.Title']
                            break
            
            if not us_title:
                if not messagebox.askyesno("Aviso", f"O Caso de Teste {test_case_id} não possui uma User Story relacionada. Deseja continuar?", icon='warning'):
                    return None, None

            return us_title, test_case_title

        except requests.RequestException as e:
            if messagebox.askyesno("Erro de API", f"Falha ao buscar relações para o CT {test_case_id}: {e}\nDeseja continuar?", icon='error'):
                return "", test_case_id
            return None, None

    def upload_attachment(self, file_path, project_name):
        """Faz o upload de um anexo para um projeto."""
        file_name = Path(file_path).name
        url = f"{ORGANIZATION_URL}/{project_name}/_apis/wit/attachments?fileName={file_name}&api-version=6.0"
        try:
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            headers = {
                'Content-Type': 'application/octet-stream',
                'Authorization': self.auth_headers['Authorization']
            }
            response = requests.post(url, headers=headers, data=file_content)
            response.raise_for_status()
            return response.json().get('url')
        except requests.RequestException as e:
            messagebox.showerror("Erro de Upload", f"Falha ao fazer upload do anexo: {e}")
            return None

    def link_attachment_to_work_item(self, work_item_id, attachment_url, project_name):
        """Associa um anexo a um item de trabalho (work item)."""
        url = f"{ORGANIZATION_URL}/{project_name}/_apis/wit/workitems/{work_item_id}?api-version=6.0"
        patch_document = [{
            "op": "add",
            "path": "/relations/-",
            "value": {
                "rel": "AttachedFile",
                "url": attachment_url,
                "attributes": {"comment": "Evidência de teste gerada por automação."}
            }
        }]
        try:
            response = requests.patch(url, headers=self.auth_headers, json=patch_document)
            response.raise_for_status()
            return True
        except requests.RequestException as e:
            messagebox.showerror("Erro de API", f"Falha ao associar anexo ao Work Item: {e}")
            return False

class TestAssistantApp:
    """Classe principal da aplicação que gerencia a UI e a lógica."""

    def __init__(self, root):
        self.root = root
        self.root.title("Assistente de Testes")
        self.root.geometry("680x500")
        try:
            self.root.iconbitmap(ICON_FILE)
        except tk.TclError:
            print(f"Aviso: Ícone '{ICON_FILE}' não encontrado.")

        self.entries = {}
        self.vars = {
            "devops_integration": tk.BooleanVar(value=False),
            "ambiente": tk.StringVar(value='HML')
        }
        
        self._setup_styles()
        self._create_widgets()

    def _setup_styles(self):
        """Configura os estilos dos widgets ttk."""
        style = ttk.Style()
        style.configure("TEntry", padding=(5, 4), relief="flat")
        style.configure("TCombobox", padding=(5, 4), relief="flat")
        style.configure("TButton", padding=(10, 5), relief="flat")
        
        style.map('TCombobox',
            fieldbackground=[('readonly', 'white'), ('focus', 'white')],
            selectbackground=[('readonly', 'white')],
            selectforeground=[('readonly', 'black')]
        )

    def _create_widgets(self):
        """Cria e posiciona todos os widgets na janela."""
        notebook = ttk.Notebook(self.root)
        notebook.pack(expand=True, fill='both', padx=10, pady=10)
        
        tab1 = ttk.Frame(notebook)
        notebook.add(tab1, text="Gerador de Documentos")

        # --- Criação dos campos de entrada ---
        for i, (key, text) in enumerate(UI_LABELS.items()):
            ttk.Label(tab1, text=text).grid(row=i, column=0, padx=10, pady=8, sticky='w')
            if key in ["env", "product"]:
                continue
            
            entry = ttk.Entry(tab1, width=40)
            entry.grid(row=i, column=1, padx=5, pady=8, sticky='we')
            self.entries[key] = entry

        # --- Botões de 'Procurar' e 'Limpar' ---
        ttk.Button(tab1, text="Procurar", command=lambda: self._browse_file(self.entries['file'])).grid(row=0, column=2)
        ttk.Button(tab1, text="Procurar", command=lambda: self._select_directory(self.entries['img_dir'])).grid(row=1, column=2)
        ttk.Button(tab1, text="Limpar Imagens", command=self._clear_images).grid(row=1, column=3, padx=5)

        # --- Dropdown de Ambiente ---
        ambiente_dropdown = ttk.Combobox(tab1, textvariable=self.vars['ambiente'], values=["DEV", "HML", "PRD"], state="readonly", width=37)
        ambiente_dropdown.grid(row=4, column=1, padx=5, pady=8, sticky='we')

        # --- Integração com DevOps ---
        devops_frame = ttk.Frame(tab1)
        devops_frame.grid(row=2, column=2, columnspan=3, sticky='w', padx=5)
        
        ttk.Checkbutton(devops_frame, text="Integrar com DevOps", variable=self.vars['devops_integration'], command=self._toggle_devops_integration).pack(side='left')
        
        try:
            key_img = ImageTk.PhotoImage(Image.open(KEY_ICON_FILE).resize((16, 16)))
            self.key_button = ttk.Button(devops_frame, image=key_img, command=self._show_key_dialog)
            self.key_button.image = key_img
            self.key_button.pack(side='left', padx=10)
        except Exception:
            ttk.Button(devops_frame, text="🔑", command=self._show_key_dialog).pack(side='left', padx=10)
            print(f"Aviso: Ícone '{KEY_ICON_FILE}' não encontrado.")

        # --- Dropdowns de Projeto ---
        self.project_dropdown = ttk.Combobox(tab1, state="disabled", width=37)
        self.project_dropdown.grid(row=5, column=1, padx=5, pady=8, sticky='we')
        
        # --- Botão Gerar e Status ---
        self.status_label = ttk.Label(tab1, text="")
        self.status_label.grid(row=9, column=0, columnspan=4, pady=10)
        
        self.generate_button = ttk.Button(tab1, text="Gerar Documento", command=self._start_generation_thread)
        self.generate_button.grid(row=8, column=1, columnspan=1, padx=5, pady=20, sticky='e')
    
    # --- Métodos de UI (Callbacks) ---

    def _browse_file(self, entry_widget):
        """Abre diálogo para selecionar arquivo .docx."""
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if file_path:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, file_path)

    def _select_directory(self, entry_widget):
        """Abre diálogo para selecionar um diretório."""
        dir_path = filedialog.askdirectory()
        if dir_path:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, dir_path)
            
    def _clear_images(self):
        """Remove arquivos de imagem do diretório selecionado."""
        image_dir_str = self.entries['img_dir'].get()
        if not image_dir_str:
            messagebox.showwarning("Atenção", "Selecione o diretório de imagens primeiro.")
            return

        image_dir = Path(image_dir_str)
        if not image_dir.is_dir():
            messagebox.showerror("Erro", f"O diretório '{image_dir}' não existe.")
            return
            
        removed_count = 0
        for file in image_dir.iterdir():
            if file.suffix.lower() in IMAGE_EXTENSIONS:
                try:
                    file.unlink()
                    removed_count += 1
                except Exception as e:
                    messagebox.showerror("Erro", f"Falha ao remover {file.name}: {e}")
        
        messagebox.showinfo("Limpeza Concluída", f"{removed_count} imagem(ns) removida(s).")

    def _toggle_devops_integration(self):
        """Ativa/desativa a integração com DevOps e atualiza a UI."""
        if self.vars['devops_integration'].get():
            self.project_dropdown.set('')
            self.project_dropdown.config(state='readonly')
            Thread(target=self._update_devops_projects, daemon=True).start()
        else:
            self.project_dropdown.set('')
            self.project_dropdown.config(state='disabled', values=[])
    
    def _update_devops_projects(self):
        """Busca projetos do DevOps e atualiza o dropdown."""
        pat = self._get_decoded_key()
        if not pat:
            self.vars['devops_integration'].set(False)
            self.project_dropdown.config(state='disabled')
            return
            
        service = DevOpsService(pat)
        projects = service.get_projects()
        
        if projects:
            self.project_dropdown.config(values=projects)
            self.status_label.config(text="Projetos carregados com sucesso.")
        else:
            messagebox.showerror("Erro", "Não foi possível carregar os projetos. Verifique sua chave de acesso.")
            self.vars['devops_integration'].set(False)
            self.project_dropdown.config(state='disabled')

    # --- Lógica de Geração ---

    def _start_generation_thread(self):
        """Inicia a geração do documento em uma thread separada para não bloquear a UI."""
        self.generate_button.config(state='disabled')
        self.status_label.config(text="Gerando documento...")
        Thread(target=self._run_generation, daemon=True).start()

    def _run_generation(self):
        """Lógica principal de criação e upload do documento."""
        try:
            # 1. Coletar e validar dados da UI
            data = {key: entry.get() for key, entry in self.entries.items()}
            data['env'] = self.vars['ambiente'].get()
            data['is_devops'] = self.vars['devops_integration'].get()
            data['project'] = self.project_dropdown.get()

            if not data['file'] or not data['img_dir']:
                raise ValueError("Os campos 'Arquivo padrão' e 'Diretório de Imagens' são obrigatórios.")
            
            image_dir = Path(data['img_dir'])
            if not any(f.suffix.lower() in IMAGE_EXTENSIONS for f in image_dir.iterdir()):
                if not messagebox.askyesno("Aviso", "Nenhuma imagem encontrada no diretório. Deseja continuar?", icon='warning'):
                    return

            # 2. Lógica de DevOps (se aplicável)
            us_title, test_case_title = "", data['test_id']
            if data['is_devops']:
                if not data['project']:
                    raise ValueError("Selecione um projeto do DevOps.")
                pat = self._get_decoded_key()
                if not pat: return

                service = DevOpsService(pat)
                us_title, test_case_title = service.get_test_case_relations(data['test_id'])
                if us_title is None: return # User cancelled

            # 3. Criar e modificar documento Word
            doc = Document(data['file'])
            
            # Inserir imagens
            self._insert_images(doc, image_dir, "Evidências")
            
            # Substituir placeholders
            placeholder_values = {
                PLACEHOLDERS['tester']: data['tester'],
                PLACEHOLDERS['test_id']: f"{data['test_id']} - {test_case_title}" if test_case_title != data['test_id'] else data['test_id'],
                PLACEHOLDERS['us']: us_title,
                PLACEHOLDERS['env']: data['env'],
                PLACEHOLDERS['profile']: data['profile'],
                PLACEHOLDERS['bugs']: data['bugs'],
                PLACEHOLDERS['result']: '☐Passed  ☒Failed' if data['bugs'] else '☒Passed  ☐Failed'
            }
            self._replace_placeholders(doc, placeholder_values)
            
            # 4. Salvar o novo documento
            new_file_name = f"{Path(data['file']).stem}_{data['test_id']}.docx"
            new_file_path = image_dir / new_file_name
            doc.save(new_file_path)

            # 5. Fazer upload para DevOps (se aplicável)
            if data['is_devops']:
                attachment_url = service.upload_attachment(new_file_path, data['project'])
                if attachment_url:
                    service.link_attachment_to_work_item(data['test_id'], attachment_url, data['project'])
                    messagebox.showinfo("Sucesso", f"Documento gerado e anexado ao DevOps!\nLocal: {new_file_path}")
                else:
                    messagebox.showwarning("Aviso", f"Documento gerado, mas falhou ao anexar ao DevOps.\nLocal: {new_file_path}")
            else:
                messagebox.showinfo("Sucesso", f"Documento gerado com sucesso!\nLocal: {new_file_path}")

        except Exception as e:
            messagebox.showerror("Erro Inesperado", f"Ocorreu um erro: {e}")
        finally:
            # Reativa o botão na thread principal
            self.root.after(0, lambda: self.generate_button.config(state='normal'))
            self.root.after(0, lambda: self.status_label.config(text=""))
    
    def _insert_images(self, doc, image_dir, marker_text):
        """Insere imagens de um diretório em um documento Word."""
        insert_paragraph = next((p for p in doc.paragraphs if marker_text in p.text), None)
        if not insert_paragraph:
            raise ValueError(f"Marcador '{marker_text}' não encontrado no documento.")
        
        # Limpa o parágrafo do marcador e prepara para inserir as imagens
        insert_paragraph.clear()
        
        image_files = sorted([f for f in image_dir.iterdir() if f.suffix.lower() in IMAGE_EXTENSIONS], key=lambda x: x.stat().st_mtime)
        
        for img_path in image_files:
            try:
                # Adiciona cada imagem em um novo parágrafo para garantir o espaçamento
                p = doc.add_paragraph()
                p.add_run().add_picture(str(img_path), width=Inches(6.5))
            except Exception as e:
                print(f"Erro ao inserir a imagem {img_path.name}: {e}")

    def _replace_placeholders(self, doc, values):
        """Substitui texto em parágrafos e tabelas do documento."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph(p, values)
        for p in doc.paragraphs:
            self._replace_in_paragraph(p, values)

    def _replace_in_paragraph(self, paragraph, values):
        """
        Função auxiliar para substituir placeholders em um único parágrafo,
        lidando com a quebra de texto em múltiplos 'runs'.
        """
        for placeholder, value in values.items():
            if placeholder in paragraph.text:
                full_text = "".join([run.text for run in paragraph.runs])
                if placeholder in full_text:
                    new_text = full_text.replace(placeholder, value)
                    
                    # Limpa os runs e cria um novo com o texto substituído
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.clear()

    # --- Gerenciamento da Chave DevOps ---

    def _get_decoded_key(self):
        """Lê e decodifica a chave do arquivo."""
        try:
            with open(KEY_FILE, "r") as f:
                encoded_key = f.read().strip()
            return base64.b64decode(encoded_key).decode('utf-8')
        except (FileNotFoundError, Exception): # Correção aqui
            messagebox.showerror("Chave não encontrada", f"Arquivo '{KEY_FILE}' não encontrado ou inválido. Por favor, cadastre sua chave.")
            return None

    def _show_key_dialog(self):
        """Mostra uma janela para o usuário inserir a chave DevOps."""
        dialog = Toplevel(self.root)
        dialog.title("Chave DevOps")
        dialog.geometry("350x150")
        dialog.transient(self.root)
        dialog.grab_set()

        ttk.Label(dialog, text="Insira seu Personal Access Token (PAT):").pack(pady=10)
        key_entry = ttk.Entry(dialog, width=50, show="*")
        key_entry.pack(pady=5, padx=10)

        def save_key():
            pat = key_entry.get().strip()
            if not pat:
                messagebox.showwarning("Erro", "A chave não pode estar vazia.", parent=dialog)
                return
            
            try:
                encoded_key = base64.b64encode(pat.encode('utf-8')).decode('utf-8')
                with open(KEY_FILE, "w") as f:
                    f.write(encoded_key)
                messagebox.showinfo("Sucesso", "Chave salva com sucesso.", parent=dialog)
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao salvar a chave: {e}", parent=dialog)
        
        ttk.Button(dialog, text="Salvar", command=save_key).pack(pady=10)
        dialog.wait_window()

if __name__ == "__main__":
    root = ThemedTk(theme="arc")
    app = TestAssistantApp(root)
    root.mainloop()